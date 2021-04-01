__doc__ = """
## Generating a poster program from formdesk output

### Author
Author: Brett G. Olivier (b.g.olivier@vu.nl)
Licence: GNU GPL v3

This program is free software: you can redistribute it and/or modify
it under the terms of the GNU General Public License as published by
the Free Software Foundation, either version 3 of the License, or
(at your option) any later version.

This program is distributed in the hope that it will be useful,
but WITHOUT ANY WARRANTY; without even the implied warranty of
MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
GNU General Public License for more details.

You should have received a copy of the GNU General Public License
along with this program.  If not, see <http://www.gnu.org/licenses/>

"""
__version__ = 0.6

# data_file = 'AIMMSannualmeeting2021_exportforBrett.xlsx'
# data_file = 'AIMMSannualmeeting2021_withoutaddressses.xlsx'
# data_file = '2103022_AIMMSannualmeeting2021_withoutaddresssesFinal.xlsx'
data_file = 'AIMMSannualmeeting2021_final_no_homeaddresses_2.xlsx'

WRITE_DB = False
WRITE_DOC = False
WRITE_GUESTLIST = True

GUEST_LIST_FILENAME = 'AIMMSday_guestlist_final_4.csv'
DB_FILE_NAME = 'aimmsday_posters.sqlite'
# DATA_COLUMNS = [
# "_fd_id",
# "_fd_seqno",
# "_fd_source",
# "_fd_status",
# "_fd_add",
# "_fd_edit",
# "first_name",
# "prefix",
# "last_name",
# "aimms_research_group",
# "job_title_1",
# "job_title1",
# "organization",
# "strees_address",
# "street_address2",
# "city",
# "zip_code",
# "country",
# "e_mail",
# "do_you_want_to_present_a_",
# "title",
# "authors",
# "short_abstract",
# "i_would_like_to_receive_a",
# ]

import os
import time
import pprint
import json
import csv
import CBNetDB
import openpyxl
import docx

assert os.environ['CONDA_DEFAULT_ENV'] == 'sandbox', 'Guess what ...'

# this sets the current year and month
CURRENT_YEAR = int(time.strftime('%Y'))
CURRENT_MONTH = int(time.strftime('%m'))

# set up current env.
ctime = time.strftime('%Y-%m-%d')
cdir = os.path.dirname(os.path.abspath(os.sys.argv[0]))
data_dir = os.path.join(cdir, '..', 'noupload', 'parse_formdesk_output', 'data')
out_dir = os.path.join(cdir, '..', 'noupload', 'parse_formdesk_output')
data_file = os.path.join(data_dir, data_file)


# set up database env.
DB_FILE_NAME = os.path.join(out_dir, DB_FILE_NAME)
DB_ACTIVE_TABLE = '\"{}\"'.format(CURRENT_YEAR)
DB_COLS = [
    'pid INT PRIMARY KEY',
    'status INT',
    'fname TEXT',
    'mname TEXT',
    'lname TEXT',
    'grp TEXT',
    'job1 TEXT',
    'job2 TEXT',
    'email TEXT',
    'poster INT',
    'title TEXT',
    'author TEXT',
    'abstract TEXT',
    'borrel INT',
    'pnumber TEXT',
    'participant INT',
]

posterDB = None


"""
json_db_file = os.path.join(out_dir, json_db_file)
# configure persistent data
if os.path.exists(json_db_file):
    F = open(json_db_file, 'r')
    DB_DATA = json.load(F)
else:
    DB_DATA = {}
    F = open(json_db_file, 'w')
    DB_DATA = json.dump(DB_DATA, F)
F.close()
"""

# open database
posterDB = CBNetDB.DBTools()
posterDB.connectSQLiteDB(DB_FILE_NAME, work_dir=out_dir)
if posterDB.getTable(DB_ACTIVE_TABLE) is None:
    posterDB.createDBTable(DB_ACTIVE_TABLE, DB_COLS)

# use excel file directly
print(data_file)
exl_wb = openpyxl.load_workbook(filename=data_file)
exl_sh = exl_wb[exl_wb.sheetnames[0]]

if WRITE_DB:
    # get max numbers
    pRegMax = len(posterDB.getColumns(DB_ACTIVE_TABLE, ['pid'])[0])
    pNumMax = posterDB.getColumns(DB_ACTIVE_TABLE, ['poster'])[0].count('1')
    print(pRegMax, pNumMax)
    for row in range(2, exl_sh.max_row + 1):
        pid = int(exl_sh['B{}'.format(row)].value)
        if not posterDB.checkEntryInColumn(DB_ACTIVE_TABLE, 'pid', pid):
            pRegMax += 1
            dta = {
                'pid': pid,
                'status': int(
                    True if exl_sh['D{}'.format(row)].value == 'Completed' else False
                ),
                'fname': exl_sh['G{}'.format(row)].value,
                'mname': exl_sh['H{}'.format(row)].value,
                'lname': exl_sh['I{}'.format(row)].value,
                'grp': exl_sh['J{}'.format(row)].value,
                'job1': exl_sh['K{}'.format(row)].value,
                'job2': exl_sh['L{}'.format(row)].value,
                'email': exl_sh['S{}'.format(row)].value,
                'poster': int(exl_sh['T{}'.format(row)].value),
                'title': exl_sh['U{}'.format(row)].value,
                'author': exl_sh['V{}'.format(row)].value,
                'abstract': exl_sh['W{}'.format(row)].value,
                'borrel': 0,  # int(exl_sh['X{}'.format(row)].value)
                'participant': pRegMax,
            }
            if dta['poster']:
                pNumMax += 1
                dta['pnumber'] = 'P{0:02d}'.format(pNumMax)
                dta['author'] = str(dta['author']).replace('\n', ',')
                dta['author'] = ', '.join([a.strip() for a in dta['author'].split(',')])

            posterDB.insertData(DB_ACTIVE_TABLE, dta, commit=False)
            print('Adding row with ID \"{}\".'.format(pid))
            # pprint.pprint(dta)
        else:
            print('Skipping existing ID \"{}\".'.format(pid))
    posterDB.commitDB()


# define text writers
def add_participant(D, table, row):
    data = posterDB.getRow(table, 'participant', row)[0]
    p = D.add_paragraph()
    if data[3] == 'None':
        p.add_run('{} {} ({})'.format(data[2], data[4], data[8]))
    else:
        p.add_run('{} {} {} ({})'.format(data[2], data[3], data[4], data[8]))
    if data[9] == '1':
        p.add_run(' - {}'.format(data[14]))
    if data[5].strip() == 'Other':
        if data[7] != 'None':
            p.add_run('\n{}'.format(data[7]))
        else:
            p.add_run('\n{}'.format(' '))
    elif data[5].strip() == 'choose one of the groups below:':
        p.add_run('\n{}'.format(' '))
    elif data[5].strip() == 'None':
        p.add_run('\n{}'.format(' '))
    else:
        p.add_run('\n{}'.format(data[5]))


def add_poster_long(D, table, row):
    data = posterDB.getRow(table, 'participant', row)[0]
    if data[9] == '1':
        p = D.add_paragraph()
        p.add_run('{}) '.format(data[14])).bold = True
        p.add_run('{}\n'.format(data[10])).bold = True
        names = str(data[11])
        names = names.replace('\n', ',')
        names = [a.strip() for a in names.split(',')]
        print(names)
        if data[3] == 'None':
            name = '{} {}'.format(data[2], data[4])
        else:
            name = '{} {} {}'.format(data[2], data[3], data[4])

        if name in names:
            r = p.add_run('{}'.format(names.pop(names.index(name))))
            r.underline = True
            r.italic = True
            r = p.add_run(', {}\n'.format(', '.join(names)))
        else:
            r = p.add_run('{}\n'.format(', '.join(names)))
        r.italic = True
        r2 = p.add_run('Contact: {} ({})\n\n'.format(data[8], data[5]))

        r3 = p.add_run('{}\n'.format(data[12]))
        r3.font.size = docx.shared.Pt(11)


def add_to_table(t, crd, max_col, table, row, ptotal):
    data = posterDB.getRow(table, 'participant', row)[0]
    alpha_map = 'FEDCBA'
    if crd[0] < 3:
        colswitch = 0
    else:
        colswitch = 1
    if crd[1] == colswitch:
        t.add_row()
        crd[0] = crd[0] + 1
        crd[1] = max_col
    if data[9] == '1':
        # get name
        if data[3] == 'None':
            name = '{} {}'.format(data[2], data[4])
        else:
            name = '{} {} {}'.format(data[2], data[3], data[4])
        # set row colour
        if crd[0] in [1, 3, 5]:
            row_colour = COLOUR_GREEN
        else:
            row_colour = COLOUR_BLUE
        print(crd[0], crd[1])
        pnumber = data[14]
        pnumber = '{}{}'.format(alpha_map[crd[1] - 1], crd[0] + 1)
        print(pnumber)
        posterDB.updateData(table, 'pid', data[0], {'pnumber': pnumber})
        cell = t.cell(crd[0], crd[1] - 1)
        #         cell.text = '{}\n{}'.format(pnumber, name)
        par = cell.paragraphs[0]
        run = par.add_run('{}\n'.format(pnumber))
        run.font.size = docx.shared.Pt(12)
        run.font.bold = True
        run.font.color.rgb = row_colour
        run = par.add_run('{}'.format(name))
        run.font.size = docx.shared.Pt(11)
        par.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
        #         run = par.runs[0]
        crd[1] = crd[1] - 1
        ptotal += 1
    return ptotal


# build document
if WRITE_DOC:

    pRegMax = len(posterDB.getColumns(DB_ACTIVE_TABLE, ['pid'])[0])
    COLOUR_BLUE = docx.shared.RGBColor(66, 36, 233)
    COLOUR_GREEN = docx.shared.RGBColor(0, 128, 0)

    Dx = docx.Document()
    _ = Dx.add_heading('AIMMS Day Poster Session v{}'.format(ctime), level=1)

    _ = Dx.add_heading('Poster session layout', level=2)
    Dx_tbl_max_col = 6
    Dx_tbl = Dx.add_table(1, Dx_tbl_max_col)
    cell_coord = [0, Dx_tbl_max_col]
    ptotal = 0
    for p_ in range(1, pRegMax + 1):
        ptotal = add_to_table(
            Dx_tbl, cell_coord, Dx_tbl_max_col, DB_ACTIVE_TABLE, p_, ptotal
        )
    posterDB.commitDB()
    Dx_tbl.style = 'Table Grid'
    Dx_tbl.autofit = True

    _ = Dx.add_heading('Poster abstracts ({})'.format(ptotal), level=2)
    for p_ in range(1, pRegMax + 1):
        add_poster_long(Dx, DB_ACTIVE_TABLE, p_)

    # Dx.add_page_break()
    #
    # _ = Dx.add_heading('Participant list ({})'.format(pRegMax), level=2)
    #
    # name_map = posterDB.getColumns(DB_ACTIVE_TABLE, ['lname', 'participant'])
    # name_map = dict(
    #     [(name_map[0][r], int(name_map[1][r])) for r in range(len(name_map[0]))]
    # )
    # names = list(name_map.keys())
    # names.sort()
    # print(name_map)
    # for name in names:
    #     add_participant(Dx, DB_ACTIVE_TABLE, name_map[name])

    posterDB.closeDB()
    Dx.save(os.path.join(out_dir, 'AIMMSday-{}.docx'.format(time.strftime('%H%M%S'))))

# write guest list
if WRITE_GUESTLIST:
    # initialize guestlist with the space manager
    guest_list = [
        ['email', 'name', ''],
        ['data@aimms.vu.nl', 'Brett Olivier', ''],
        ['b.g.olivier@vu.nl', 'Brett Olivier', ''],
    ]

    # add lastminute guests
    guest_list.append(['mj.smit@vu.nl', 'Martine Smit', ''])
    guest_list.append(['jeff@createscapes.com', 'Jeff Povlo', ''])
    guest_list.append(['p.van.hoorn@vu.nl', 'Peter van Hoorn', ''])
    guest_list.append(['s.m.anbuhl@vu.nl', 'Stephanie Anbuhl', ''])
    guest_list.append(['f.m.paulussen@vu.nl', 'F.M. Paulussen', ''])
    guest_list.append(['a.e.idemudia@student.vu.nl', 'Alison Idemudia', ''])
    guest_list.append(['timo.hamers@vu.nl', 'Timo Hamers', ''])
    guest_list.append(['t.a.hagel@student.vu.nl', 'Thomas Hagel', ''])

    # load details from excel spreadsheet
    for row in range(2, exl_sh.max_row + 1):
        if exl_sh['H{}'.format(row)].value is None:
            name = '{} {}'.format(
                exl_sh['G{}'.format(row)].value, exl_sh['I{}'.format(row)].value
            )
        else:
            name = '{} {} {}'.format(
                exl_sh['G{}'.format(row)].value,
                exl_sh['H{}'.format(row)].value,
                exl_sh['I{}'.format(row)].value,
            )
        guest_list.append([exl_sh['S{}'.format(row)].value, name, ''])

    with open(os.path.join(out_dir, GUEST_LIST_FILENAME), 'w', newline='') as F:
        csvw = csv.writer(F)
        csvw.writerows(guest_list)

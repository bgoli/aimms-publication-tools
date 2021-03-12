__doc__ = """
## Generating AIMMS manuscript reports

There is a two stage process to create the AIMMS paper report

### Generate a custom report in PURE
You need to create an Excel report in PURE. To do this you need elevated user rights to the AIMMS contex and the correct report format. Get this from  Brett).

### Parse data and update database
This code use a sql database to store publication reports. There are two things that are important here, the first being the normal run and the second database maintenance.

#### Normal operation
This is what you do on a regular basis, first copy the Excel spreadsheet you generated from PURE into the `\data` directory. Next add a line of code that loads that file.

```python
data_file = 'AIMMS_research_2021-8_02_21.xls'
```

Save the file and run this script from a terminal
```bash
python aimms_monthly_articles.py
````

This will update the database and JSON index files.

#### Database update (each new year)
The database *aimmsDB.sqlite* is set up to use the table *publications* for the current years results. At the begining of a calendar year simply rename
the *publications* table to the previous year, for example, 2020. In this way a history of AIMMS publications is maintained.

### Report generation

Next you can generate the report using

```bash
python generate_publication_report.py
```
This will result in two Word documents:

```bash
AIMMS_publication_report-2021-02-08.docx
AIMMS_publications_for_newsletter-2021-02-08.docx
```
A detailed list of new publications and a shortened version suitable for publication in the newsletter.


### Author
Author: Brett G. Olivier (b.g.olivier@vu.nl)
Licence: GNU GPL v3

This program is free software: you can redistribute it and/or modify
it under the terms of the GNU General Public License as published by
the Free Software Foundation, either version 3 of the License, or
(at your option) any later version.

This program is distributed in the hope that it will be useful,
but WITHOUT ANY WARRANTY; without even the implied warranty of
MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
GNU General Public License for more details.

You should have received a copy of the GNU General Public License
along with this program.  If not, see <http://www.gnu.org/licenses/>

"""
__version__ = 1.3

import os
import time

# import pprint
import CBNetDB
import docx
import json

# import numpy

assert os.environ['CONDA_DEFAULT_ENV'] == 'sandbox', 'Guess what ...'

# override date range and treat all papers as processed
DO_ALL = False

# set current month
CURRENT_MONTH = int(time.strftime('%m'))
CURRENT_YEAR = int(time.strftime('%Y'))
# manual override for specific month/year
# CURRENT_YEAR = 2020
# CURRENT_MONTH = 12

ctime = time.strftime('%Y-%m-%d')
cdir = os.path.dirname(os.path.abspath(os.sys.argv[0]))
data_dir = os.path.join(cdir, 'data')

# set up database env.
# default DB name: 'aimmsDB.sqlite'
DB_FILE_NAME = 'aimmsDB.sqlite'
# default table name: 'publications'
DB_ACTIVE_TABLE = 'publications'

# load "old format" json file
"""
json_file = os.path.join(data_dir, 'current_doi_dataset.json')
with open(json_file, 'r') as F:
    parsed_data = json.load(F)
"""
aimmsDB = CBNetDB.DBTools()
aimmsDB.connectSQLiteDB(DB_FILE_NAME, work_dir=cdir)

sldata = aimmsDB.getColumns(DB_ACTIVE_TABLE, ['doi', 'month', 'newdata'])

new_papers = []
month_papers = []
other_new_papers = []
print(len(sldata[0]))
for d in range(len(sldata[0])):
    print(int(sldata[1][d]), CURRENT_MONTH, eval(sldata[2][d]))
    if DO_ALL:
        month_papers.append(aimmsDB.getRow(DB_ACTIVE_TABLE, 'doi', sldata[0][d])[0])
    elif int(sldata[1][d]) in (CURRENT_MONTH, CURRENT_MONTH - 1) and eval(sldata[2][d]):
        new_papers.append(aimmsDB.getRow(DB_ACTIVE_TABLE, 'doi', sldata[0][d])[0])
    elif int(sldata[1][d]) in (CURRENT_MONTH, CURRENT_MONTH - 1):
        month_papers.append(aimmsDB.getRow(DB_ACTIVE_TABLE, 'doi', sldata[0][d])[0])
    elif int(sldata[1][d]) in range(1, CURRENT_MONTH + 1) and eval(sldata[2][d]):
        other_new_papers.append(aimmsDB.getRow(DB_ACTIVE_TABLE, 'doi', sldata[0][d])[0])

aimmsDB.closeDB()

del sldata

# pprint.pprint(new_papers)

Dx = docx.Document()
Dx2news = docx.Document()

_ = Dx.add_heading('AIMMS publication report for: {}'.format(ctime), level=1)
_ = Dx2news.add_heading('AIMMS publication report for: {}'.format(ctime), level=1)

# add new papers to report index
for d in new_papers:
    # print(d)
    p0 = Dx.add_paragraph(
        '{} ({}-{})'.format(d[6][:60], d[4], d[2]), style='List Number'
    )

# add already processed monthly papers to report
for d in month_papers:
    # print(d)
    p0 = Dx.add_paragraph(style='List Number')
    p0.add_run('{} ({}-{})'.format(d[6][:60], d[4], d[2])).italic = True

# add new out of date scope papers to report
for d in other_new_papers:
    # print(d)
    p0 = Dx.add_paragraph(style='List Number')
    p0.add_run('{} ({}-{})'.format(d[6][:60], d[4], d[2])).italic = True


def add_detail_list(doc, D):
    x = doc.add_paragraph(D[7], style='List Bullet')
    x = doc.add_paragraph(D[9], style='List Bullet')
    x = doc.add_paragraph(D[10], style='List Bullet')
    x = doc.add_paragraph(D[0], style='List Bullet')
    x = doc.add_paragraph('Corresponding author: {}'.format(D[8]), style='List Bullet')
    x = doc.add_paragraph(
        'Published {} (early online {})'.format(D[3], D[5]), style='List Bullet'
    )
    x = doc.add_paragraph('Processed: {}-{}'.format(D[4], D[2]), style='List Bullet')
    del x


def add_newsletter_item(doc, D):
    p = doc.add_paragraph()
    p.add_run(D[7] + ' ')
    p.add_run(D[6] + ' ').bold = True
    p.add_run('({}, {})[{}]'.format(D[10], D[3], D[0]))


# add already processed papers to the general report
_ = Dx2news.add_heading(level=3).add_run(
    'New papers: {}-{}/{}'.format(CURRENT_YEAR, CURRENT_MONTH - 1, CURRENT_MONTH)
)

# add already processed papers to the weekly report
"""
_ = Dx2news.add_heading(level=3).add_run(
    'Processed papers: {}-{}/{}'.format(CURRENT_YEAR, CURRENT_MONTH - 1, CURRENT_MONTH)
)
"""

cntr = 1

for d in new_papers:
    # detailed doc
    h = Dx.add_heading(level=3)
    h.add_run('{}) {}'.format(cntr, d[6]))
    add_detail_list(Dx, d)
    if DO_ALL:
        p = Dx.add_paragraph(d[11].replace(d[6], '')[:301] + ' ...')
    else:
        p = Dx.add_paragraph(d[11].replace(d[6], ''))
    p.add_run(' ')

    # newsletter
    add_newsletter_item(Dx2news, d)

    cntr += 1

for d in month_papers:
    # detailed doc
    h = Dx.add_heading(level=3)
    h.add_run('{}) {}'.format(cntr, d[6])).italic = True
    add_detail_list(Dx, d)
    p = Dx.add_paragraph(d[11].replace(d[6], '')[:200] + ' ...')
    p.add_run(' ')

    # newsletter
    # add_newsletter_item(Dx2news, d)

    cntr += 1

for d in other_new_papers:
    # detailed doc
    h = Dx.add_heading(level=3)
    h.add_run('{}) {}'.format(cntr, d[6])).italic = True
    add_detail_list(Dx, d)
    p = Dx.add_paragraph(d[11].replace(d[6], '')[:200] + ' ...')
    p.add_run(' ')

    # newsletter
    # add_newsletter_item(Dx2news, d)

    cntr += 1


Dx.save('AIMMS_publication_report-{}.docx'.format(ctime))
Dx2news.save('AIMMS_publications_for_newsletter-{}.docx'.format(ctime))
time.sleep(2)
